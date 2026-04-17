"""Generate tiny PDF and DOCX fixtures from the Markdown SRS.

Run once: `uv run python tests/fixtures/make_fixtures.py`.
The generated .pdf and .docx are committed so tests don't need WeasyPrint at runtime.
"""

from pathlib import Path

from docx import Document

HERE = Path(__file__).parent
MD_PATH = HERE / "srs_small.md"
DOCX_PATH = HERE / "srs_small.docx"
PDF_PATH = HERE / "srs_small.pdf"


def build_docx() -> None:
    doc = Document()
    for line in MD_PATH.read_text(encoding="utf-8").splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.strip():
            doc.add_paragraph(line)
    doc.save(DOCX_PATH)


def build_pdf() -> None:
    from markdown import markdown
    from weasyprint import HTML

    html = markdown(MD_PATH.read_text(encoding="utf-8"))
    HTML(string=html).write_pdf(PDF_PATH)


if __name__ == "__main__":
    build_docx()
    print(f"wrote {DOCX_PATH}")
    try:
        build_pdf()
        print(f"wrote {PDF_PATH}")
    except Exception as exc:  # noqa: BLE001
        print(f"skipped PDF: {exc}")
