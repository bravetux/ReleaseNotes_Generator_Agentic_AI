"""Multi-format SRS loader tool (PDF, DOCX, TXT, MD)."""

import logging
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from strands import tool

from release_notes_agent.config.settings import settings

logger = logging.getLogger(__name__)

SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".md"}
MAX_CHARS = settings.max_srs_chars_per_file


def _read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


_READERS = {
    ".pdf": _read_pdf,
    ".docx": _read_docx,
    ".txt": _read_text,
    ".md": _read_text,
}


def _expand(paths: str | list[str]) -> list[Path]:
    items = [paths] if isinstance(paths, str) else list(paths)
    out: list[Path] = []
    for raw in items:
        p = Path(raw)
        if p.is_dir():
            out.extend(
                sorted(c for c in p.iterdir() if c.suffix.lower() in SUPPORTED_EXTS)
            )
        else:
            out.append(p)
    return out


def _load_one(path: Path) -> dict:
    ext = path.suffix.lower()
    reader = _READERS.get(ext)
    if reader is None:
        return {
            "filename": path.name,
            "format": ext.lstrip("."),
            "error": f"unsupported extension {ext}",
        }
    try:
        text = reader(path)
    except Exception as exc:  # noqa: BLE001
        logger.warning("failed to read %s: %s", path, exc)
        return {
            "filename": path.name,
            "format": ext.lstrip("."),
            "error": str(exc),
        }
    truncated = False
    from release_notes_agent.tools import srs_loaders as _self

    if len(text) > _self.MAX_CHARS:
        text = text[: _self.MAX_CHARS]
        truncated = True
        logger.warning("truncated %s to %d chars", path.name, _self.MAX_CHARS)
    return {
        "filename": path.name,
        "format": ext.lstrip("."),
        "raw_text": text,
        "truncated": truncated,
    }


@tool
def load_srs(paths: str | list[str]) -> list[dict]:
    """Load SRS files from one or many paths (files or folders).

    Args:
        paths: A single path or a list of paths. Paths may be files
            (.pdf/.docx/.txt/.md) or folders. Folders are scanned non-recursively
            for supported extensions.

    Returns:
        A list of dicts, one per file. Each has ``filename``, ``format``,
        and either ``raw_text`` + ``truncated`` on success or ``error`` on failure.
    """
    return [_load_one(p) for p in _expand(paths)]
